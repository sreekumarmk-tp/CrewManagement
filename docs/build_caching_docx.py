"""
Generate docs/CACHING_DESIGN.docx — a styled Word version of the caching plan
(Steps 1-4) with rendered flow-diagram images.

Run:  python docs/build_caching_docx.py
Deps: python-docx, matplotlib, Pillow  (already installed in this env)

Diagrams are drawn with matplotlib (boxes + arrows) into docs/_diagrams/*.png,
then embedded into the Word document. Re-run to regenerate after edits.
"""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from math import hypot

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
DIAG = os.path.join(HERE, "_diagrams")
os.makedirs(DIAG, exist_ok=True)

# ── Diagram drawing ──────────────────────────────────────────────────────────
PALETTE = {
    "frontend": ("#DBEAFE", "#2563EB"),
    "backend":  ("#DCFCE7", "#15803D"),
    "llm":      ("#EDE9FE", "#7C3AED"),
    "db":       ("#FFEDD5", "#C2410C"),
    "cache":    ("#FEE2E2", "#DC2626"),
    "decision": ("#FEF9C3", "#A16207"),
    "neutral":  ("#F1F5F9", "#475569"),
    "win":      ("#D1FAE5", "#059669"),
}


def draw(nodes, edges, fname, xlim, ylim):
    w = (xlim[1] - xlim[0]) / 12.0
    h = (ylim[1] - ylim[0]) / 12.0
    fig, ax = plt.subplots(figsize=(w, h), dpi=200)
    ax.set_xlim(*xlim)
    ax.set_ylim(*ylim)
    ax.set_aspect("equal")
    ax.axis("off")

    patches = {}
    for nid, n in nodes.items():
        face, edge = PALETTE[n.get("kind", "neutral")]
        x, y = n["xy"]
        bw, bh = n["w"], n["h"]
        box = FancyBboxPatch(
            (x - bw / 2, y - bh / 2), bw, bh,
            boxstyle="round,pad=0.3,rounding_size=1.4",
            linewidth=1.7, edgecolor=edge, facecolor=face, zorder=2,
        )
        ax.add_patch(box)
        patches[nid] = box
        ax.text(x, y, n["text"], ha="center", va="center",
                fontsize=n.get("fs", 9), color="#0f172a", zorder=3,
                linespacing=1.25)

    for e in edges:
        src, dst = e[0], e[1]
        label = e[2] if len(e) > 2 else ""
        rad = e[3] if len(e) > 3 else 0.0
        lpos = e[4] if len(e) > 4 else None  # explicit (x, y) label position
        ca, cb = nodes[src]["xy"], nodes[dst]["xy"]
        arr = FancyArrowPatch(
            ca, cb, patchA=patches[src], patchB=patches[dst],
            arrowstyle="-|>", mutation_scale=15, linewidth=1.5,
            color="#334155", connectionstyle=f"arc3,rad={rad}",
            shrinkA=3, shrinkB=3, zorder=1,
        )
        ax.add_patch(arr)
        if label:
            if lpos is not None:
                lx, ly = lpos
            else:
                mx, my = (ca[0] + cb[0]) / 2, (ca[1] + cb[1]) / 2
                dx, dy = cb[0] - ca[0], cb[1] - ca[1]
                dist = hypot(dx, dy) or 1
                nx, ny = -dy / dist, dx / dist
                off = rad * dist * 0.5
                lx, ly = mx + nx * off, my + ny * off
                if rad == 0:
                    ly += 2.2
            ax.text(lx, ly, label, ha="center", va="center", fontsize=7.6,
                    color="#1e293b", zorder=4,
                    bbox=dict(boxstyle="round,pad=0.25", fc="white",
                              ec="#cbd5e1", lw=0.6, alpha=0.95))

    path = os.path.join(DIAG, fname)
    fig.savefig(path, bbox_inches="tight", pad_inches=0.12, facecolor="white")
    plt.close(fig)
    return path


def build_diagrams():
    paths = {}

    # A — baseline flow (no caching)
    paths["baseline"] = draw(
        {
            "UI":   {"xy": (16, 74), "w": 26, "h": 15, "kind": "frontend", "text": "Next.js UI\n(Zustand store)"},
            "API":  {"xy": (57, 74), "w": 22, "h": 15, "kind": "backend", "text": "FastAPI\nroutes"},
            "REPO": {"xy": (100, 74), "w": 28, "h": 15, "kind": "neutral", "text": "crew_repository\nSELECT * FROM crew"},
            "PG":   {"xy": (100, 42), "w": 24, "h": 13, "kind": "db", "text": "PostgreSQL"},
            "MA":   {"xy": (57, 42), "w": 24, "h": 15, "kind": "llm", "text": "Managed Agents\ncoordinator session"},
            "SPEC": {"xy": (57, 13), "w": 44, "h": 15, "kind": "llm", "text": "4 specialist sub-agents\ncrew_matching · travel ·\nnotification · compliance"},
        },
        [
            ("UI", "API", "GET /crew/*  ·  POST /workflow/*"),
            ("API", "REPO", "every call"),
            ("REPO", "PG", ""),
            ("API", "MA", "run_turn"),
            ("MA", "SPEC", ""),
            ("SPEC", "REPO", "tool calls", -0.35, (74, 30)),
        ],
        "01_baseline.png", (0, 120), (0, 90),
    )

    # B — four steps at a glance
    paths["overview"] = draw(
        {
            "S1": {"xy": (32, 86), "w": 50, "h": 13, "kind": "llm", "text": "Step 1 ✓ — LLM layer\nMaximize Managed-Agents prompt caching"},
            "W1": {"xy": (98, 86), "w": 38, "h": 11, "kind": "win", "text": "↓ input tokens billed"},
            "S2": {"xy": (32, 63), "w": 50, "h": 13, "kind": "neutral", "text": "Step 2 — Config layer\nlru_cache for repeated file reads"},
            "W2": {"xy": (98, 63), "w": 38, "h": 11, "kind": "win", "text": "↓ disk reads"},
            "S3": {"xy": (32, 40), "w": 50, "h": 13, "kind": "cache", "text": "Step 3 — Data layer\nRedis cache-aside for crew queries"},
            "W3": {"xy": (98, 40), "w": 38, "h": 11, "kind": "win", "text": "↓ DB load & latency"},
            "S4": {"xy": (32, 17), "w": 50, "h": 13, "kind": "frontend", "text": "Step 4 — Frontend layer\nSWR + HTTP cache headers"},
            "W4": {"xy": (98, 17), "w": 38, "h": 11, "kind": "win", "text": "↓ redundant API calls"},
        },
        [("S1", "W1", ""), ("S2", "W2", ""), ("S3", "W3", ""), ("S4", "W4", "")],
        "02_overview.png", (0, 122), (5, 98),
    )

    # C — Step 1 prompt-cache hit/miss
    paths["step1"] = draw(
        {
            "APP":  {"xy": (24, 78), "w": 34, "h": 14, "kind": "backend", "text": "run_turn()\nsend user.message\n(dynamic tail only)", "fs": 8.5},
            "MA":   {"xy": (74, 78), "w": 28, "h": 13, "kind": "llm", "text": "Managed Agents\n(server)"},
            "DEC":  {"xy": (74, 52), "w": 34, "h": 15, "kind": "decision", "text": "static prefix unchanged?\n(system prompt + context)", "fs": 8.5},
            "HIT":  {"xy": (38, 28), "w": 32, "h": 14, "kind": "win", "text": "HIT — reuse at\n~10% token cost"},
            "MISS": {"xy": (98, 28), "w": 32, "h": 14, "kind": "cache", "text": "MISS — process full\nprefix, store copy"},
            "OUT":  {"xy": (68, 8), "w": 52, "h": 12, "kind": "neutral", "text": "events + usage  (input, output, cache_read, cache_creation)", "fs": 8},
        },
        [
            ("APP", "MA", "send"),
            ("MA", "DEC", ""),
            ("DEC", "HIT", "warm", -0.2),
            ("DEC", "MISS", "cold / first call", 0.2),
            ("HIT", "OUT", ""),
            ("MISS", "OUT", ""),
        ],
        "03_step1.png", (0, 122), (0, 92),
    )

    # D — Step 2 lru_cache (two runtime-hot file sources + the excluded one)
    paths["step2"] = draw(
        {
            # Lane A — skills.json (registry.py)
            "A1": {"xy": (20, 92), "w": 34, "h": 11, "kind": "neutral", "text": "_custom_skill_refs()"},
            "A2": {"xy": (20, 78), "w": 34, "h": 11, "kind": "neutral", "text": "custom_skill_id_to_name()"},
            "HA": {"xy": (64, 85), "w": 38, "h": 16, "kind": "decision", "text": "_load_skills_cache()\n@lru_cache(maxsize=1)\nskills.json · registry.py", "fs": 8},
            "DA": {"xy": (113, 93), "w": 32, "h": 11, "kind": "db", "text": "read + parse\nskills.json"},
            "MA": {"xy": (113, 76), "w": 32, "h": 11, "kind": "win", "text": "cached dict\n(no disk I/O)"},
            # Lane B — role/skill markdown (loader.py)
            "B1": {"xy": (20, 48), "w": 34, "h": 11, "kind": "backend", "text": "build_instructions(agent)"},
            "B2": {"xy": (20, 34), "w": 34, "h": 11, "kind": "backend", "text": "list_skill_files(agent)"},
            "HB": {"xy": (64, 41), "w": 38, "h": 16, "kind": "decision", "text": "@lru_cache\n(keyed by agent)\nrole/skill .md · loader.py", "fs": 8},
            "DB": {"xy": (113, 49), "w": 36, "h": 12, "kind": "db", "text": "read system_prompt.md\n+ skill / shared .md"},
            "MB": {"xy": (113, 33), "w": 32, "h": 11, "kind": "win", "text": "cached str / stems\n(no disk I/O)"},
            # Excluded
            "EX": {"xy": (70, 9), "w": 128, "h": 13, "kind": "neutral", "fs": 7.6,
                   "text": "NOT cached: policy SKILL.md folders (crew-travel-policy, visa-and-transit-requirements,\n"
                           "port-clearance-procedures, repatriation-rules) — read only by the offline\n"
                           "scripts.attach_skills upload; runtime caching for these is server-side (Step 1)."},
        },
        [
            ("A1", "HA", ""),
            ("A2", "HA", ""),
            ("HA", "DA", "first call", 0.0, (90, 91)),
            ("HA", "MA", "later calls", 0.0, (90, 79)),
            ("B1", "HB", ""),
            ("B2", "HB", ""),
            ("HB", "DB", "first call", 0.0, (90, 47)),
            ("HB", "MB", "later calls", 0.0, (90, 35)),
        ],
        "04_step2.png", (0, 142), (0, 104),
    )

    # E — Step 3 Redis cache-aside
    paths["step3"] = draw(
        {
            "REQ":   {"xy": (18, 84), "w": 30, "h": 12, "kind": "frontend", "text": "GET /crew/sign-on"},
            "SVC":   {"xy": (18, 58), "w": 32, "h": 12, "kind": "backend", "text": "get_sign_on_crew()"},
            "CHK":   {"xy": (56, 58), "w": 30, "h": 14, "kind": "decision", "text": "Redis GET\ncrew:signon"},
            "RET":   {"xy": (104, 84), "w": 28, "h": 12, "kind": "win", "text": "return list"},
            "DB":    {"xy": (104, 56), "w": 32, "h": 12, "kind": "db", "text": "SELECT … pool='signon'"},
            "STORE": {"xy": (82, 30), "w": 32, "h": 12, "kind": "cache", "text": "Redis SETEX\nTTL 30 min"},
            "UPD":   {"xy": (18, 16), "w": 32, "h": 12, "kind": "backend", "text": "update_crew() commit"},
            "DEL":   {"xy": (56, 12), "w": 32, "h": 12, "kind": "cache", "text": "Redis DEL\ncrew:signon / signoff"},
        },
        [
            ("REQ", "SVC", ""),
            ("SVC", "CHK", ""),
            ("CHK", "RET", "HIT", -0.15),
            ("CHK", "DB", "MISS"),
            ("DB", "STORE", ""),
            ("STORE", "RET", ""),
            ("UPD", "DEL", "invalidate"),
            ("DEL", "CHK", "next read\nrepopulates", -0.3, (40, 35)),
        ],
        "05_step3.png", (0, 128), (0, 96),
    )

    # F — Step 4 frontend SWR
    paths["step4"] = draw(
        {
            "COMP":  {"xy": (18, 72), "w": 32, "h": 14, "kind": "frontend", "text": "React components\n(crew table, pickers)"},
            "SWR":   {"xy": (58, 72), "w": 30, "h": 14, "kind": "decision", "text": "SWR / React Query\ncache"},
            "FAST":  {"xy": (100, 86), "w": 30, "h": 12, "kind": "win", "text": "render instantly\n(no network)"},
            "FETCH": {"xy": (100, 56), "w": 30, "h": 12, "kind": "frontend", "text": "axios GET /crew/*"},
            "API":   {"xy": (100, 28), "w": 34, "h": 13, "kind": "backend", "text": "FastAPI\nCache-Control: max-age=300", "fs": 8},
            "REDIS": {"xy": (58, 16), "w": 30, "h": 12, "kind": "cache", "text": "Step 3 Redis\ncache-aside"},
            "PG":    {"xy": (18, 16), "w": 24, "h": 12, "kind": "db", "text": "PostgreSQL"},
        },
        [
            ("COMP", "SWR", ""),
            ("SWR", "FAST", "fresh", -0.18),
            ("SWR", "FETCH", "stale / empty"),
            ("FETCH", "API", ""),
            ("API", "REDIS", ""),
            ("REDIS", "PG", ""),
            ("FETCH", "SWR", "revalidate", 0.35),
        ],
        "06_step4.png", (0, 122), (5, 96),
    )

    # G — implementation order
    paths["order"] = draw(
        {
            "S2": {"xy": (18, 50), "w": 26, "h": 20, "kind": "neutral", "text": "Step 2 ✓\nlru_cache\nfile reads"},
            "S1": {"xy": (50, 50), "w": 26, "h": 20, "kind": "llm", "text": "Step 1 ✓\nprompt-cache\n+ metrics"},
            "S3": {"xy": (82, 50), "w": 26, "h": 20, "kind": "cache", "text": "Step 3\nRedis crew\ncache-aside"},
            "S4": {"xy": (114, 50), "w": 26, "h": 20, "kind": "frontend", "text": "Step 4\nfrontend SWR\n+ headers"},
        },
        [("S2", "S1", ""), ("S1", "S3", ""), ("S3", "S4", "")],
        "07_order.png", (0, 132), (30, 70),
    )

    return paths


# ── Word document ────────────────────────────────────────────────────────────
INK = RGBColor(0x0F, 0x17, 0x2A)
MUTE = RGBColor(0x47, 0x55, 0x69)
ACCENT = RGBColor(0x25, 0x63, 0xEB)


def style_base(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10.5)
    n.font.color.rgb = INK


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = ACCENT if level <= 1 else INK
    return p


def para(doc, text, italic=False, color=None, size=None, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(after)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            lead, rest = it
            r = p.add_run(lead)
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(it)


def image(doc, path, width=6.3, caption=None):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = MUTE


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    return t


def acceptance(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("Acceptance signal — ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(10)


def build_doc(paths):
    doc = Document()
    style_base(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Title block
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Caching Design & Implementation Plan")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("CrewManagement — Maritime Crew Orchestrator")
    rs.font.size = Pt(12)
    rs.font.color.rgb = MUTE
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run("Branch: Work_Cacheing   ·   Scope: Steps 1–4   ·   Step 5 deferred   ·   Design only")
    rm.font.size = Pt(9.5)
    rm.font.color.rgb = MUTE
    doc.add_paragraph()

    # 1. Context
    h(doc, "1. Context — what the system looks like today", 1)
    para(doc, "Two of the cheapest wins exist because Redis is already provisioned in "
              "docker-compose but never instantiated, and axios is already centralized in one "
              "file. Most of this plan is “use what's there”, not “add new infrastructure”.")
    table(doc,
          ["Layer", "Stack", "Caching today"],
          [
              ["Frontend", "Next.js 15 + React 19 + Zustand, bare axios", "None — crew lists fetched every mount"],
              ["Backend API", "FastAPI (async)", "None — no cache layer or headers"],
              ["LLM layer", "Anthropic Managed Agents (sessions)", "Session reused across Phase 1→2 only"],
              ["Database", "PostgreSQL 16 (async SQLAlchemy)", "None — full-table scans every call"],
              ["Infra", "Docker Compose; Redis 7 declared", "redis_url in config, never used"],
          ])
    doc.add_paragraph()
    h(doc, "End-to-end flow today (no caching)", 3)
    para(doc, "Every component on the right is hit fresh on every request. The four steps insert "
              "a cache in front of each hot path.", color=MUTE, size=9.5)
    image(doc, paths["baseline"], caption="Figure 1 — Current request flow with no caching")

    # 2. Overview
    h(doc, "2. The four steps at a glance", 1)
    image(doc, paths["overview"], caption="Figure 2 — Each step and the cost it removes")
    table(doc,
          ["Step", "Layer", "Effort", "Payoff", "New dependency?"],
          [
              ["1  (done)", "LLM (Managed Agents)", "Low", "High (token cost + latency)", "No"],
              ["2  (done)", "Backend config + skill files", "Trivial", "Small–Medium", "No (functools stdlib)"],
              ["3", "Backend data", "Medium", "Medium (DB load, ~50–100 ms)", "Redis client (infra exists)"],
              ["4", "Frontend", "Medium", "Medium (UX, fewer calls)", "SWR / TanStack Query (npm)"],
          ])

    # Step 1
    doc.add_page_break()
    h(doc, "Step 1 — Maximize Managed-Agents server-side prompt caching  [Implemented]", 1)
    h(doc, "What it does", 2)
    para(doc, "Reduces the input tokens billed on every coordinator/specialist turn. The large, "
              "static system prompts (COORDINATOR_SYSTEM_ROLE and each specialist's system_prompt()) "
              "are re-processed by the model on every turn. Prompt caching lets the platform reuse a "
              "stored copy of that static prefix at ~10% of the normal input-token cost, and also "
              "lowers latency.")
    h(doc, "How it does it — important nuance", 2)
    para(doc, "This project does NOT call the raw Messages API, so there is no place to hand-insert "
              "cache_control blocks per request. It uses Managed Agents (client.beta.sessions in "
              "backend/agents/managed/client.py), where prompt caching is platform-managed. Our job "
              "is to enable and keep the cache warm, not to build it:")
    bullets(doc, [
        ("Keep system prompts byte-stable. ", "The cache key is the prompt prefix. Any change to "
         "COORDINATOR_SYSTEM_ROLE (registry.py:48) or a specialist's system_prompt() — even "
         "whitespace — invalidates the cache and the next turn pays full price."),
        ("Put dynamic content at the end. ", "In _run_turn() (client.py:385) the per-turn message "
         "is sent as a user.message. Keep variable per-workflow data (crew id, reason, phase data) "
         "at the tail so it doesn't disturb the cached prefix."),
        ("Reuse the session across phases — already done. ", "A single coordinator session spans "
         "Phase 1 and Phase 2; keeping it means accumulated context stays cache-resident between the "
         "two human-separated turns. Preserve this."),
        ("Verify the cache is hit — done. ", "_run_turn() now captures cache_read_input_tokens / "
         "cache_creation_input_tokens from each span.model_request_end. They accrue onto WorkflowState "
         "(cache_read_tokens / cache_creation_tokens) via _record_usage, which also refines total_cost "
         "(reads ~0.1x, writes ~1.25x of input rate). get_metrics() aggregates them and computes "
         "cache_hit_rate = reads / (reads + writes) x 100. /monitoring/metrics returns those fields and "
         "/monitoring/workflows/active exposes per-workflow cache tokens; the dashboard shows a Cache "
         "Hit Rate KPI backed by the new SystemMetrics fields."),
    ])
    para(doc, "Implemented with the claude-api skill, which confirmed the authoritative model_usage "
              "shape (input/output/cache_read/cache_creation) and that Managed Agents caches the prompt "
              "prefix automatically — there is no cache_control to insert.",
         italic=True, color=MUTE, size=9.5)
    image(doc, paths["step1"], caption="Figure 3 — Prompt-cache hit/miss per turn")
    acceptance(doc, "Met. /monitoring/metrics returns cache_hit_rate (with cache_read_tokens / "
                    "cache_creation_tokens) and the dashboard shows a Cache Hit Rate KPI. As the "
                    "byte-stable prompts persist, cache_read_tokens climbs and the ratio rises while "
                    "billed input tokens per workflow drop. Metric math verified (zero-safe).")

    # Step 2
    doc.add_page_break()
    h(doc, "Step 2 — In-memory lru_cache for repeated config & skill-file reads  [Implemented]", 1)
    h(doc, "What it does", 2)
    para(doc, "Eliminates repeated disk reads of the two file sources that are re-read on a "
              "request/runtime hot path:")
    bullets(doc, [
        ("backend/skills.json. ", "Re-opened and re-parsed every time a specialist config is built "
         "(_custom_skill_refs and custom_skill_id_to_name in registry.py)."),
        ("Role/skill markdown under backend/agents/skills/<agent>/. ", "build_instructions() reads "
         "system_prompt.md plus every skill *.md and shared/*.md and concatenates them. It runs "
         "inside CrewMatchingAgent.__init__ (crew_matching_agent.py:75), and a specialist is "
         "constructed fresh on every workflow phase — so each phase previously re-read a handful of "
         "files. list_skill_files() similarly re-globs the folder for the monitoring API (loader.py)."),
    ])
    para(doc, "All of these files are effectively immutable at runtime (they change only when a "
              "developer edits them, after which the server restarts), so reading each once per "
              "process is enough.")
    h(doc, "What it deliberately does NOT cache", 2)
    para(doc, "The policy SKILL.md folders (crew-travel-policy, visa-and-transit-requirements, "
              "port-clearance-procedures, repatriation-rules) are read only by _skill_files() -> "
              "upload_skill() -> the offline one-shot scripts/attach_skills.py. They are never read "
              "on a request path — the hosted Anthropic agent loads them server-side. An lru_cache "
              "only pays off across repeated calls in a long-lived process, so caching a one-shot "
              "upload reader saves nothing and would even be counter-productive in a dev edit -> "
              "re-upload loop. Their runtime caching belongs to Step 1, not local disk caching.")
    h(doc, "How it does it", 2)
    bullets(doc, [
        ("skills.json. ", "_load_skills_cache() with @lru_cache(maxsize=1); both call sites read "
         "through it. The cached dict is treated as read-only."),
        ("Role/skill markdown. ", "build_instructions() is decorated with @lru_cache(maxsize=None) "
         "(keyed by agent name); list_skill_files() is backed by a cached _skill_file_stems() that "
         "returns an immutable tuple, and the public function returns a fresh list(...) so callers "
         "may mutate it freely."),
        ("Failure handling. ", "lru_cache does not cache exceptions, so a missing file on an early "
         "call is simply retried on the next one rather than wedging an empty result."),
    ])
    image(doc, paths["step2"], width=6.6,
          caption="Figure 4 — Two cached readers (skills.json, loader markdown); policy folders excluded")
    acceptance(doc, "skills.json and each agent's role/skill markdown are read from disk exactly "
                    "once per process. Verified: build_instructions.cache_info() reports 1 miss + N "
                    "hits, and mutating a list_skill_files() result does not corrupt the cache.")

    # Step 3
    doc.add_page_break()
    h(doc, "Step 3 — Redis cache-aside for crew-list queries", 1)
    h(doc, "What it does", 2)
    para(doc, "Caches the results of get_sign_on_crew() and get_sign_off_crew() "
              "(crew_repository.py:15–30). These run full-table SELECTs and are called on every "
              "frontend page load and at every workflow initiation, even though the crew pool changes "
              "infrequently. A short-TTL cache removes most of that DB load and shaves request latency.")
    h(doc, "How it does it — cache-aside pattern", 2)
    para(doc, "A new backend/services/cache_service.py instantiates an async Redis client against the "
              "existing settings.redis_url (config.py:36). The repository functions become cache-aware:")
    bullets(doc, [
        ("Read. ", "Build a key (crew:signon / crew:signoff). Try Redis first. HIT → deserialize and "
         "return, never touch Postgres. MISS → run the existing SQLAlchemy query, store in Redis with "
         "a TTL (suggested 30 min), return it."),
        ("Write / invalidation. ", "update_crew() (crew_repository.py:45) already mutates pool/status. "
         "After a successful commit it must DELETE the affected keys so the next read repopulates "
         "fresh data — correctness over TTL alone."),
        ("Graceful degradation. ", "If Redis is unreachable, fall through to the DB query. The cache "
         "is an optimization, never a hard dependency."),
    ])
    para(doc, "get_crew_by_id() is a point lookup used during compliance; leave it uncached initially, "
              "or cache per-id with the same invalidation hook. Start with the two list queries — that "
              "is where the repeated full scans are.", italic=True, color=MUTE, size=9.5)
    image(doc, paths["step3"], caption="Figure 5 — Cache-aside read path and invalidation on write")
    acceptance(doc, "Repeated /crew/* calls within the TTL issue zero SQL; a crew mutation immediately "
                    "invalidates the cache; a Redis outage falls back to DB with no errors.")

    # Step 4
    doc.add_page_break()
    h(doc, "Step 4 — Frontend data caching (SWR / React Query + HTTP headers)", 1)
    h(doc, "What it does", 2)
    para(doc, "Stops the UI from re-fetching slow-changing crew lists redundantly and gives it "
              "stale-while-revalidate behavior: instant render from cache, silent background refresh. "
              "Today the crew lists are fetched with bare axios in frontend/src/lib/api.ts and parked "
              "in Zustand with no dedup, no revalidation, and total loss on refresh.")
    h(doc, "How it does it", 2)
    bullets(doc, [
        ("Client cache library. ", "Introduce SWR (or TanStack Query) for the read-only crew endpoints "
         "— crewApi.getSignOnCrew() / getSignOffCrew() (api.ts:14–15). It provides request dedup, "
         "caching with revalidation, and background refresh, replacing the manual mount-time fetch."),
        ("What stays uncached. ", "POST /workflow/* mutations and the WebSocket live event stream are "
         "inherently dynamic — leave them as-is. SWR wraps only the GET reads."),
        ("HTTP cache headers (backend side). ", "Add Cache-Control: public, max-age=300 to the GET "
         "crew endpoints via FastAPI so the browser/axios layer can also short-circuit. This pairs "
         "with Step 3: Redis cuts DB load, HTTP headers cut network round-trips."),
    ])
    image(doc, paths["step4"], caption="Figure 6 — SWR cache in front of the cached backend")
    acceptance(doc, "Navigating between views does not re-issue crew fetches within the freshness "
                    "window; multiple components mounting at once trigger a single network request; "
                    "data still refreshes on revalidation.")

    # Order + cross-cutting
    doc.add_page_break()
    h(doc, "3. Recommended implementation order", 1)
    image(doc, paths["order"], caption="Figure 7 — Suggested sequencing")
    para(doc, "Step 2 is a safe warm-up. Step 1 adds the usage metrics that prove later wins. Step 3 "
              "must land before Step 4 so the frontend caches data that is itself already DB-cheap. "
              "Step 4 closes the loop at the edge.")

    h(doc, "4. Cross-cutting concerns", 1)
    bullets(doc, [
        ("Correctness over hit-rate. ", "Every cache needs an invalidation story (Step 3's "
         "update_crew hook is the critical one). A stale crew list shown during a live sign-off is "
         "worse than a slow one."),
        ("Graceful degradation. ", "Redis (Step 3) and the prompt cache (Step 1) are optimizations — "
         "the app must work identically if they are cold or unavailable."),
        ("Observability. ", "Step 1's cache-hit metric and a Redis hit/miss counter (Step 3) should "
         "surface on the existing /monitoring endpoints so wins are measurable, not assumed."),
        ("Out of scope (Step 5, deferred). ", "Moving the in-memory StateService._workflows dict to "
         "Redis for horizontal scaling. Revisit only when running more than one backend instance."),
    ])

    out = os.environ.get("CACHING_DOCX_OUT") or os.path.join(HERE, "CACHING_DESIGN.docx")
    doc.save(out)
    return out


if __name__ == "__main__":
    paths = build_diagrams()
    out = build_doc(paths)
    print("Wrote", out)
