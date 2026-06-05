"""Minimal Markdown -> .docx converter for the L1 SignalFabric plan.

Handles the subset of Markdown used in PLAN.md: headings, paragraphs, bullet/checkbox
lists, GitHub tables, fenced code blocks, bold (**x**), and inline code (`x`).
Links [text](url) are rendered as their text.
"""
import os
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC, OUT = sys.argv[1], sys.argv[2]
SRC_DIR = os.path.dirname(os.path.abspath(SRC))

# Max printable width inside default 1-inch margins on Letter (8.5") = 6.5"
MAX_W_IN = 6.3

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_runs(paragraph, text):
    """Add inline-formatted runs (bold, code, links) to a paragraph."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif part.startswith("["):
            m = LINK.match(part)
            paragraph.add_run(m.group(1) if m else part)
        else:
            paragraph.add_run(part)


def render_table(lines):
    rows = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    # drop the |---|---| separator row
    body = [r for i, r in enumerate(rows) if not (i == 1 and set("".join(r)) <= set("-: "))]
    ncols = max(len(r) for r in body)
    t = doc.add_table(rows=0, cols=ncols)
    t.style = "Light Grid Accent 1"
    for i, r in enumerate(body):
        cells = t.add_row().cells
        for j in range(ncols):
            txt = r[j] if j < len(r) else ""
            p = cells[j].paragraphs[0]
            add_runs(p, txt)
            if i == 0:
                for run in p.runs:
                    run.bold = True


lines = open(SRC, encoding="utf-8").read().splitlines()
i = 0
n = len(lines)
while i < n:
    line = lines[i]

    # fenced code block
    if line.startswith("```"):
        i += 1
        buf = []
        while i < n and not lines[i].startswith("```"):
            buf.append(lines[i])
            i += 1
        i += 1
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(12)
        r = p.add_run("\n".join(buf))
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        continue

    # standalone image:  ![alt](path)
    img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
    if img:
        alt, path = img.group(1), img.group(2)
        abspath = path if os.path.isabs(path) else os.path.join(SRC_DIR, path)
        if os.path.exists(abspath):
            from PIL import Image as _Img
            w_px, h_px = _Img.open(abspath).size
            # scale to fit width, preserving aspect ratio
            width_in = min(MAX_W_IN, w_px / 150.0)  # PNGs rendered at 150 dpi
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(abspath, width=Inches(width_in))
            if alt:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap.add_run(alt)
                r.italic = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        else:
            doc.add_paragraph(f"[missing image: {path}]")
        i += 1
        continue

    # horizontal rule
    if line.strip() == "---":
        doc.add_paragraph().add_run("─" * 40).font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        i += 1
        continue

    # table block
    if line.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
        block = []
        while i < n and lines[i].strip().startswith("|"):
            block.append(lines[i])
            i += 1
        render_table(block)
        doc.add_paragraph()
        continue

    # headings
    m = re.match(r"^(#{1,6})\s+(.*)", line)
    if m:
        level = len(m.group(1))
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", m.group(2))
        text = LINK.sub(r"\1", text)
        text = text.replace("`", "")
        if level == 1:
            h = doc.add_heading(text, level=0)
        else:
            h = doc.add_heading(level=min(level, 4))
            add_runs(h, m.group(2))
        i += 1
        continue

    # blockquote
    if line.startswith(">"):
        p = doc.add_paragraph(style="Intense Quote")
        add_runs(p, line.lstrip("> ").rstrip())
        i += 1
        continue

    # list items (bullets, checkboxes)
    lm = re.match(r"^(\s*)[-*]\s+(.*)", line)
    if lm:
        content = lm.group(2)
        checkbox = ""
        cb = re.match(r"^\[([ xX])\]\s+(.*)", content)
        if cb:
            checkbox = "☑ " if cb.group(1).lower() == "x" else "☐ "
            content = cb.group(2)
        p = doc.add_paragraph(style="List Bullet")
        if checkbox:
            p.add_run(checkbox)
        add_runs(p, content)
        i += 1
        continue

    # blank
    if not line.strip():
        i += 1
        continue

    # plain paragraph
    p = doc.add_paragraph()
    add_runs(p, line)
    i += 1

doc.save(OUT)
print("wrote", OUT)
