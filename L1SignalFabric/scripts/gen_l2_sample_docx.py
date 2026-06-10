"""Generate `docs/L2_Connector_Sample_Data.docx` — the narrative counterpart to
the workbook. Same sample data (imported from gen_l2_sample_workbook), rendered
as printable tables with the computed L2 result per row.

Run:  python scripts/gen_l2_sample_docx.py
"""
from __future__ import annotations

import json

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

from core.signal import SourceSystem
import scripts.gen_l2_sample_workbook as g

NAVY = RGBColor(0x1F, 0x4E, 0x78)
GREEN = RGBColor(0x37, 0x56, 0x23)
GREY = RGBColor(0x59, 0x59, 0x59)
HEAD_BG = "1F4E78"
L2_BG = "E2EFDA"


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_fill})
    tcPr.append(shd)


def set_text(cell, text, *, bold=False, color=None, size=8.5, mono=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if mono:
        run.font.name = "Consolas"
    if color is not None:
        run.font.color.rgb = color


def add_table(doc, headers, rows, widths, l2_last=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_text(c, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=8.5)
        shade(c, HEAD_BG)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            is_l2 = l2_last and i == len(row) - 1
            if isinstance(v, list):
                v = "; ".join(v)
            set_text(cells[i], v, size=8 if not is_l2 else 8,
                     color=GREEN if is_l2 else None, bold=is_l2)
            if is_l2:
                shade(cells[i], L2_BG)
    for i, w in enumerate(widths):
        for r in table.rows:
            r.cells[i].width = Inches(w)
    return table


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def para(doc, text, *, italic=False, color=None, size=9.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p


def connector_section(doc, title, subtitle, cols, rows, widths):
    heading(doc, title, level=2)
    para(doc, subtitle, italic=True, color=GREY, size=9)
    headers = cols + ["→ Resulting L2 record(s)"]
    body = []
    for d, ev in rows:
        body.append([d[c] for c in cols] + [g.facet_summary(ev)])
    add_table(doc, headers, body, widths)
    doc.add_paragraph()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(0.5))
    usable = 10.0  # ~ landscape Letter width minus margins

    # ---- title page ----
    t = doc.add_heading("L2 Connector Sample Data", level=0)
    for r in t.runs:
        r.font.color.rgb = NAVY
    para(doc, "10 human-style records per connector → the unified L2 Record (Org / Entity / Ops facets).",
         italic=True, color=GREY, size=11)
    para(doc, "")
    para(doc, "Every '→ Resulting L2 record(s)' value is COMPUTED by running the sample through "
              "l2.record.project_record — it is the actual L2 output, not a hand-written guess.", size=10)
    para(doc, "One SignalEvent → one L2Record envelope (provenance) + a list of FACETS, one per map:", size=10)
    for b in ("OrgMap facet — node / edge / signoff_event (tribal-knowledge graph)",
              "EntityMap facet — crew / contract / vessel / port record to MERGE",
              "OpsMap facet — a process-mining event-log row (per crew-change case)"):
        doc.add_paragraph(b, style="List Bullet")
    para(doc, "Contracts: EntityMap = L1_TO_L2_ENTITY_EVENT_TRIGGERS.md · OpsMap = "
              "L1_TO_L2_INGESTION_CONTRACT.md · OrgMap = L1_TO_L2_ORGMAP_CONTRACT.md "
              "(+ l2/orgmap.py tribal edges). Envelope: L1_TO_L2_UNIFIED_RECORD.md.", size=9, color=GREY)
    doc.add_page_break()

    # ---- connector sections ----
    w = usable
    c, r = g.crewdb_rows()
    connector_section(doc, "ERP · Crew DB  (CREW_DB)",
                      "Crew master rows → EntityMap crew triggers + an OrgMap Crew node. Covers upserted / signed_on / signed_off / deleted.",
                      ["crew_id", "pool", "name", "rank", "nationality", "vessel", "port", "certifications", "op"],
                      r, [0.7, 0.5, 1.1, 0.9, 0.7, 1.1, 0.7, 1.5, 0.4, 2.4])
    c, r = g.contract_rows()
    connector_section(doc, "ERP · Contract / CLM  (CONTRACT_CLM)",
                      "Engagement contract rows → EntityMap contract.upserted + an OrgMap Contract node.",
                      ["contract_id", "crew_name", "rank", "vessel", "port", "start_date", "status"],
                      r, [1.0, 1.2, 0.9, 1.5, 0.7, 0.8, 0.7, 3.2])
    c, r = g.vesselport_rows()
    connector_section(doc, "ERP · Vessel / Port DB  (VESSEL_PORT_DB)",
                      "Reference rows → EntityMap vessel.upserted / port.upserted (a 'type' of port makes a Port node, else a Vessel).",
                      ["id", "type", "name", "imo_or_locode", "fleet", "country", "status"],
                      r, [0.6, 0.5, 1.5, 1.0, 1.3, 1.0, 0.9, 3.2])
    c, r = g.slack_rows()
    connector_section(doc, "Slack  (SLACK)",
                      "Messages / reactions / joins → OrgMap edges. A sign-on/off notice also yields an Entity trigger; with a workflow_id, an Ops event.",
                      ["entity", "channel", "user", "text", "workflow_id"],
                      r, [0.8, 1.0, 0.9, 3.6, 0.8, 2.9])
    c, r = g.email_rows(SourceSystem.GMAIL)
    connector_section(doc, "Gmail  (GMAIL, metadata only)",
                      "From/To/Subject → OrgMap EMAILED edge. l2Intent = CREATE_SIGNOFF_EVENT → SignOffEvent node; workflow_id → Ops event; clean 'Name (Rank)' subject → Entity trigger.",
                      ["from", "to", "subject", "l2Intent", "workflow_id"],
                      r, [1.5, 1.6, 2.6, 1.4, 0.7, 2.2])
    c, r = g.email_rows(SourceSystem.OUTLOOK)
    connector_section(doc, "Outlook  (OUTLOOK, metadata only)",
                      "Identical shape to Gmail (Microsoft Graph mail) — same OrgMap / SignOffEvent / Ops projection.",
                      ["from", "to", "subject", "l2Intent", "workflow_id"],
                      r, [1.5, 1.6, 2.6, 1.4, 0.7, 2.2])
    c, r = g.notion_rows()
    connector_section(doc, "Notion  (NOTION)",
                      "Pages / databases → OrgMap document nodes (label = entity type).",
                      ["entity", "title", "url", "author", "last_edited"],
                      r, [0.8, 3.0, 2.2, 1.1, 0.9, 2.0])
    c, r = g.sharepoint_rows()
    connector_section(doc, "SharePoint  (SHAREPOINT)",
                      "Drive items / list items → OrgMap document nodes.",
                      ["entity", "name", "site", "url", "modified_by"],
                      r, [0.8, 2.4, 1.3, 2.4, 1.1, 2.0])
    c, r = g.database_rows()
    connector_section(doc, "Database  (DATABASE, generic SQL CDC / outbox)",
                      "Outbox rows → OrgMap entity nodes keyed by (table, pk). op carries INSERT / UPDATE / DELETE.",
                      ["table", "op", "pk", "changed_fields"],
                      r, [1.5, 0.7, 1.0, 4.0, 2.8])

    # ---- OpsMap events ----
    heading(doc, "OpsMap · Workflow Event Log  (POST /opsmap/events)", level=2)
    para(doc, "Three complete cases covering every event_type and every agent_name. case_id = workflow_id is the join key; L2 sorts by timestamp and maps each event to the Activity shown.",
         italic=True, color=GREY, size=9)
    ops_body = []
    for case_id, et, an, ts, data in g.OPS_EVENTS:
        activity = {
            "workflow_created": "Sign-Off Initiated", "crew_updated": "Sign-Off Confirmed",
            "auto_compliance": "Compliance Check", "sign_on_initiated": "Compliance Check (manual)",
            "crew_signed_on": "Signed On  [TERMINAL ✓]", "sign_on_rejected": "Sign-On Rejected  [TERMINAL ✗]",
            "workflow_failed": "Workflow Failed  [TERMINAL ✗]",
            "agent_completed": {"Crew Matching Agent": "Crew Matching", "Travel Agent": "Travel Arranged",
                                "Notification Agent": "Crew Notified", "Compliance Agent": "Compliance Check"}.get(an, "?"),
        }[et]
        ops_body.append([case_id, et, an, ts, json.dumps(data), activity])
    add_table(doc, ["case_id", "event_type", "agent_name", "timestamp", "data (JSON)", "→ L2 Activity"],
              ops_body, [0.7, 1.3, 1.3, 1.6, 2.9, 2.2])
    doc.add_paragraph()

    # ---- OrgMap structure & manning ----
    heading(doc, "OrgMap · Ownership Structure  (POST /orgmap/structure)", level=2)
    para(doc, "Company → Fleet → Vessel ownership. Vessel names MUST match EntityMap Vessel.name exactly (join key).",
         italic=True, color=GREY, size=9)
    add_table(doc, ["company", "fleet", "vessels (must match Vessel.name)", "→ L2 nodes & edges"],
              [[co, fl, ve, f"Company:{co} -OWNS-> Fleet:{fl} -OPERATES-> {ve.count(';')+1} vessel(s)"]
               for co, fl, ve in g.ORG_STRUCTURE], [2.3, 1.9, 3.0, 2.8])
    doc.add_paragraph()

    heading(doc, "OrgMap · Manning Scale  (POST /orgmap/manning)", level=2)
    para(doc, "Required headcount per rank. rank names MUST match crew rank (the join between 'required' and 'have'). 'default' applies to every vessel; a vessel name is a per-ship override.",
         italic=True, color=GREY, size=9)
    add_table(doc, ["scope", "rank (must match crew rank)", "required", "→ L2 edge"],
              [[sc, rk, str(rq), f"-REQUIRES_RANK {{required:{rq}}}-> Rank:{rk}"] for sc, rk, rq in g.MANNING],
              [2.3, 2.6, 0.9, 4.2])
    doc.add_page_break()

    # ---- Full envelope JSON ----
    heading(doc, "Full L2Record Envelope — representative record per connector", level=2)
    para(doc, "The complete wire shape L1 emits: provenance header + facets[] fan-out — exactly what project_record returns.",
         italic=True, color=GREY, size=9)
    reps = [
        ("ERP – CrewDB", g.crewdb_rows()[1][0][1]),
        ("Slack (sign-on notice + workflow_id)", g.slack_rows()[1][1][1]),
        ("Gmail (sign-off, l2Intent)", g.email_rows(SourceSystem.GMAIL)[1][0][1]),
        ("Database (generic outbox row)", g.database_rows()[1][0][1]),
    ]
    for label, ev in reps:
        rec = g.project_record(ev)
        heading(doc, label, level=3)
        p = doc.add_paragraph()
        run = p.add_run(json.dumps(rec.model_dump(mode="json"), indent=2, ensure_ascii=False))
        run.font.name = "Consolas"
        run.font.size = Pt(8)
    doc.add_page_break()

    # ---- Coverage matrix ----
    heading(doc, "Coverage Matrix — every L2 record type ↔ its sample", level=2)
    coverage = [
        ("EntityMap", "crew.upserted", "ERP – CrewDB", "SNO-1000 Juan dela Cruz (signon, Available)"),
        ("EntityMap", "crew.signed_on", "ERP – CrewDB", "SOF-2000 Miguel Torres (signon, assigned)"),
        ("EntityMap", "crew.signed_off", "ERP – CrewDB", "SOF-3002 Oleksandr Tkachenko (signoff)"),
        ("EntityMap", "crew.deleted", "ERP – CrewDB", "SNO-1015 Jose Mendoza (op = DELETE)"),
        ("EntityMap", "contract.upserted", "ERP – Contract", "CT-SOF-2000 (Miguel Torres engagement)"),
        ("EntityMap", "vessel.upserted", "ERP – Vessel-Port", "VS-001 MV Pacific Star (type = vessel)"),
        ("EntityMap", "port.upserted", "ERP – Vessel-Port", "PT-SIN Singapore (type = port)"),
        ("OpsMap", "workflow_created", "OpsMap – Workflow", "wf-3001 event 1"),
        ("OpsMap", "agent_completed (×4 agents)", "OpsMap – Workflow", "wf-3001 / wf-3002"),
        ("OpsMap", "crew_updated", "OpsMap – Workflow", "wf-3001 event 5"),
        ("OpsMap", "auto_compliance", "OpsMap – Workflow", "wf-3001 event 6"),
        ("OpsMap", "sign_on_initiated", "OpsMap – Workflow", "wf-3002 event 1"),
        ("OpsMap", "crew_signed_on (terminal ✓)", "OpsMap – Workflow", "wf-3001 event 7"),
        ("OpsMap", "sign_on_rejected (terminal ✗)", "OpsMap – Workflow", "wf-3002 event 3"),
        ("OpsMap", "workflow_failed (terminal ✗)", "OpsMap – Workflow", "wf-3003 event 1"),
        ("OrgMap (tribal)", "edge POSTED_IN / REACTED_IN / MEMBER_OF", "Slack", "message / reaction / join"),
        ("OrgMap (tribal)", "edge EMAILED", "Gmail / Outlook", "all e-mail rows"),
        ("OrgMap (tribal)", "node SignOffEvent", "Gmail / Outlook", "rows with l2Intent"),
        ("OrgMap (tribal)", "node (Document / generic entity)", "Notion / SharePoint / Database", "all rows"),
        ("OrgMap (structure)", "Company / Fleet + OWNS / OPERATES", "OrgMap – Structure", "2 companies, 4 fleets"),
        ("OrgMap (structure)", "Rank + REQUIRES_RANK {required}", "OrgMap – Manning", "default + override"),
        ("OrgMap (structure)", "Crew -HAS_RANK-> Rank (derived)", "ERP – CrewDB", "derived from crew 'rank'"),
    ]
    add_table(doc, ["L2 Map", "L2 record / facet type", "Produced by", "Example sample"],
              [list(x) for x in coverage], [1.5, 3.3, 1.8, 3.4], l2_last=False)

    out = "docs/L2_Connector_Sample_Data.docx"
    doc.save(out)
    print(f"saved {out}")


if __name__ == "__main__":
    build()
