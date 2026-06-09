"""
Generator: a narrative Word document explaining the layered architecture and the
L3 Intelligence Graph workflow (input -> processing -> output -> where it goes).

Run:  python -m scripts.build_story_docx
Out:  docs/Maritime_Crew_Orchestrator_Workflow.docx
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parents[2] / "docs" / "Maritime_Crew_Orchestrator_Workflow.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
CODE_FILL = "F4F4F4"
NOTE_FILL = "FDF3D7"
INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _rich(p, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        else:
            p.add_run(part)


def para(doc, text="", style=None):
    p = doc.add_paragraph(style=style)
    _rich(p, text)
    return p


def bullet(doc, text):
    return para(doc, text, style="List Bullet")


def numbered(doc, text):
    return para(doc, text, style="List Number")


def code(doc, text):
    for line in text.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        _shade(p, CODE_FILL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def note(doc, text):
    p = doc.add_paragraph()
    _rich(p, text)
    p.paragraph_format.left_indent = Pt(8)
    _shade(p, NOTE_FILL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, hh in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        _rich(cell.paragraphs[0], hh)
        for r in cell.paragraphs[0].runs:
            r.bold = True
        _shade(cell.paragraphs[0], "D9E2F3")
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(row):
            cells[i].paragraphs[0].text = ""
            _rich(cells[i].paragraphs[0], c)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def h(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level <= 1:
        run.font.color.rgb = ACCENT
    return heading


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_heading(level=0)
    title.add_run("Maritime Crew Orchestrator").font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sr = sub.add_run("Layered Architecture & the L3 Intelligence Graph Workflow")
    sr.italic = True
    sr.font.size = Pt(12)
    para(doc, "How a crew sign-off becomes a ranked, explained replacement shortlist: "
              "the L1-L4 story, and an in-depth look at what the Intelligence Graph layer "
              "takes in, how it processes it, what it outputs, and where that output goes.")

    # ── 1. Business story ──────────────────────────────────────────────────────
    h(doc, "1. The Business Story", 1)
    para(doc, "A **Chief Officer on an LNG vessel signs off unexpectedly in Singapore**. "
              "Operations must identify a qualified replacement within hours while ensuring "
              "certifications, contract rules, and vessel requirements are all met.")
    para(doc, "Done manually this takes hours of phone calls, spreadsheet checks and "
              "document chasing. The orchestrator turns it into a single click: the system "
              "ingests the sign-off event, gathers the facts it needs, has specialist agents "
              "evaluate every available candidate, ranks them with a transparent rationale, "
              "signs on the best fit, and notifies the operators — in seconds.")
    table(doc, ["Audience", "The question this document answers"], [
        ["Business", "How did we find the replacement, and how much faster is it?"],
        ["Operations", "Who is the pick, who are the fallbacks, who do I contact?"],
        ["Engineering", "What is the input, how is it processed, where does the output go?"],
        ["QA", "Why was a candidate selected, and why was another ranked lower / rejected?"],
    ])

    # ── 2. Architecture at a glance ────────────────────────────────────────────
    h(doc, "2. The Layered Architecture at a Glance", 1)
    para(doc, "The platform is a pipeline of four layers. Each consumes the layer before it "
              "and produces a cleaner, more decision-ready artifact:")
    code(doc, """
 L1 SignalFabric      L2 Knowledge Graph        L3 Intelligence Graph        L4 Decision Graph
 (event streams)  ->  (facts: who/what/rules) ->  (reasoning: ranked      ->  (records the
 sign-off event      EntityMap/OpsMap/OrgMap       shortlist + rationale)      decision + precedent)
                                                          |
                                                          v
                                                    Operators notified
""")
    table(doc, ["Layer", "Role", "Produces", "In this codebase"], [
        ["**L1 SignalFabric**", "Ingest & normalize the sign-off event; create a workflow context",
         "A workflow context (rank, port, vessel…)", "The sign-off trigger + `SignOffContext`"],
        ["**L2 Knowledge Graph**", "Hold the facts needed to decide — candidates, certs, contracts, port rules",
         "Entities + relationships (graph facts)", "Apache AGE `maritime` graph / rule data"],
        ["**L3 Intelligence Graph**", "Reason over the facts: evaluate every candidate, rank, explain, notify",
         "A ranked, explained shortlist + notifications", "`agents/intelligence/` (this document's focus)"],
        ["**L4 Decision Graph**", "Record the decision + build precedent history",
         "Persisted decision / precedent", "Not built here — a separate team owns it"],
    ])
    note(doc, "**Status note:** L4 is intentionally **not** implemented in this codebase. "
              "L3 produces exactly the artifact L4 would consume.")

    # ── 3. The supporting layers ───────────────────────────────────────────────
    h(doc, "3. The Supporting Layers (L1, L2, L4)", 1)

    h(doc, "L1 — SignalFabric (the event arrives)", 2)
    para(doc, "The process starts when a **sign-off event** arrives (from Crew Management, "
              "ERP, email, or Operations). L1 normalizes it and creates the **workflow "
              "context** — in code, a `SignOffContext` carrying the vacated rank, grade, "
              "vessel, port, sign-off date, and contract period. This context is the unit of "
              "work that flows into L3.")
    code(doc, """SignOffContext(
    vacated_rank = "Chief Officer",
    vacated_grade = "Grade A",
    vessel = "MV Pacific Star",
    port = "Singapore",
    contract_period_months = 6,
)""")

    h(doc, "L2 — Knowledge Graph (the facts)", 2)
    para(doc, "L2 holds the **facts** L3 reasons over, as a property graph (Apache AGE, "
              "queried with Cypher) co-located in the same PostgreSQL instance. Representative "
              "facts and edges:")
    code(doc, """(:Seafarer)-[:NATIONAL_OF]->(:Country)
(:Seafarer)-[:HOLDS]->(:Certificate)
(:Seafarer)-[:ASSIGNED_TO]->(:Vessel)
(:Port)-[:RESTRICTS]->(:Country)        # nationality restrictions at a port
(:Port/:Rank)-[:REQUIRES]->(:Certificate)""")
    para(doc, "L3 reads L2 through a single seam, `agents/intelligence/graph_gateway.py`, "
              "which works in two modes set by `GRAPH_BACKEND`:")
    bullet(doc, "**`fallback`** (default): the same facts from Python rule data — demoable "
                "with no graph infrastructure.")
    bullet(doc, "**`age`**: live openCypher queries against the L2 `maritime` graph. In this "
                "mode L3 also **sources its candidate pool from the graph's `(:Seafarer)` "
                "nodes** instead of the relational table.")
    note(doc, "Either way the shape returned to L3 is identical, so the investigators never "
              "need to know which backend ran. Switching to the real graph is a config flip.")

    h(doc, "L4 — Decision Graph (downstream)", 2)
    para(doc, "L4 would **record** L3's ranked result as an auditable decision and build a "
              "precedent index (e.g. “CO placements at Singapore scoring > 80”) that "
              "could later feed back into L3's ranking. It is owned by a separate team and is "
              "not part of this codebase; L3's output is already the exact trace it needs.")

    # ── 4. L3 deep dive ────────────────────────────────────────────────────────
    h(doc, "4. The Intelligence Graph Layer (L3) — In Depth", 1)
    para(doc, "L3 is the **reasoning layer**. Its one job: turn a sign-off vacancy into a "
              "**ranked, explained shortlist of replacement crew**, and notify the operators "
              "who act on it. It is built as a **Supervisor + 3 specialist investigators** "
              "pattern. The rest of this section follows the data: **input → processing "
              "→ output → where the output goes.**")

    h(doc, "4.1 INPUT — what L3 takes in, and where it comes from", 2)
    para(doc, "L3 receives two things:")
    table(doc, ["Input", "What it is", "Where it is taken from"], [
        ["**The vacancy**", "The `SignOffContext` (rank, grade, vessel, port, period)",
         "From L1 — derived from the **departing crew member's own record** when triggered "
         "by “Initiate Sign Off”, or from an explicit rank+port for an ad-hoc match"],
        ["**The candidate pool**", "All crew available for sign-on (the people to rank)",
         "The relational `crew` table (`get_sign_on_crew`) by default; or the **L2 graph's "
         "`(:Seafarer)` nodes** when `GRAPH_BACKEND=age`"],
        ["**L2 facts (per-run)**", "Port nationality restrictions, rank-required certs",
         "`graph_gateway` → Cypher over L2 (age) or rule data (fallback)"],
    ])
    para(doc, "Entry points (HTTP API, mounted at `/api/v1/intelligence`):")
    code(doc, """POST /intelligence/match          { crew_id }          # vacancy from a departing crew member
POST /intelligence/match-context  { vacated_rank, port } # explicit vacancy
POST /intelligence/sign-on        { crew_id }            # place the chosen #1""")

    h(doc, "4.2 PROCESSING — how the input becomes a ranking", 2)
    para(doc, "The `IntelligenceSupervisor` orchestrates one run. It is a pure orchestrator: "
              "it never scores candidates itself.")
    numbered(doc, "**Load the candidate pool** (relational table or L2 graph).")
    numbered(doc, "**Fan out to all 3 investigators in parallel** (`asyncio.gather`). Each "
                  "investigator independently evaluates **every** candidate on its own "
                  "dimension and returns a per-candidate score (0–1) + an eligibility "
                  "(hard-gate) flag + reasons + the facts it consulted.")
    numbered(doc, "**Fuse → rank** (`ranking.fuse`): drop anyone a hard gate disqualified, "
                  "then blend the surviving per-dimension scores into one fused score and sort.")
    numbered(doc, "**Build the fit graph** (`build_fit_graph`): a node/edge projection of the "
                  "run for visualization.")
    numbered(doc, "**Notify operators** (`OperatorNotifier`) on the correct channel.")
    numbered(doc, "**Stream `intel_*` events** over the WebSocket throughout, so the UI "
                  "animates the run live.")
    para(doc, "The three investigators and what each checks:")
    table(doc, ["Investigator (weight)", "Hard gates (disqualify)", "Soft signals (score 0–1)"], [
        ["**Crew Intel** (50%)", "Unavailable; rank wrong-family or > 1 step off the ladder",
         "Exact rank +0.55 / adjacent +0.35; STCW valid +0.25; base certs +0.10; experience up to +0.10"],
        ["**Vessel Ops** (30%)", "Missing a vessel-mandated cert; port-restricted nationality "
         "without a valid visa (L2 `Port-[:RESTRICTS]->Country`)",
         "Mandated certs +0.45; meets sea-time +0.25; already at join port +0.30 / can relocate +0.18"],
        ["**Contract/Wage** (20%)", "None (advisory — never blocks alone)",
         "Wage in band +0.6 (below +0.5, over scaled); contract period in envelope +0.4 (MLC max +0.25)"],
    ])
    para(doc, "**Fusion & the decision (who ranks top):** a candidate is disqualified if **any** "
              "investigator marks it ineligible. Survivors are scored by a weighted blend and "
              "sorted — the #1 is simply the highest fused score. It is **deterministic** "
              "(stable tiebreak by crew_id), so the same inputs always give the same ranking.")
    code(doc, """score = (0.50 * crew + 0.30 * vessel + 0.20 * contract) * 100
ranked.sort(key = -score, then crew_id)      # top-N returned""")
    para(doc, "**Worked example** (real output, CO @ Singapore):")
    code(doc, """#1 Juan dela Cruz   crew 0.96  vessel 1.00  contract 1.00
   = 0.50*0.96 + 0.30*1.00 + 0.20*1.00 = 0.48 + 0.30 + 0.20 = 0.98  -> 98.0

#2 Piotr Kowalski   loses ~3.6 pts on Vessel Ops (must relocate; Juan was
   already at the join port)                                          -> 94.7

(17 of 23 candidates were disqualified by hard gates: missing GMDSS,
 unavailable, wrong rank family, restricted nationality, etc.)""")

    h(doc, "4.3 OUTPUT — what L3 produces", 2)
    para(doc, "Every run returns one `IntelResult` object (also the live WebSocket payloads). "
              "Its fields:")
    table(doc, ["Field", "What it carries"], [
        ["`status`", "`matched` | `no_crew_found` | `error`"],
        ["`candidates[]`", "The ranked top-N: rank_position, name, fused **score**, **rationale**, "
         "and **dimension_scores** (crew/vessel/contract)"],
        ["`reports[]`", "Per-investigator detail for **every** candidate: each dimension's score, "
         "eligibility, **reasons**, **signals** (the facts consulted) and **applied** rules — "
         "this powers the explainability"],
        ["`fit_graph`", "The derived node/edge graph (Vacancy → Candidate → Dimension → "
         "L2-Fact) + the `backend` label (age/fallback)"],
        ["`notifications[]`", "Who was notified, on which channel, and delivery status"],
        ["`pool_size`, `disqualified`", "How many were assessed and filtered"],
        ["`timing`", "first_event_ms / total_ms (the SLO measurements)"],
    ])

    h(doc, "4.4 WHERE THE OUTPUT GOES", 2)
    para(doc, "The `IntelResult` fans out to four consumers:")
    bullet(doc, "**The operator UI (primary).** The HTTP response populates the frontend store; "
                "the **Shortlist tab** lists the top selected candidates, the **explainability** "
                "panel shows why #1 was selected (reasons + sources) and why each fallback scored "
                "lower, and the **fit graph** renders the reasoning visually. Live `intel_*` "
                "events animate the run as it happens.")
    bullet(doc, "**The agent's sign-on action.** The rank-1 candidate is signed on "
                "(`/intelligence/sign-on`) — moved into the onboard pool — closing the loop.")
    bullet(doc, "**Operators**, via the `OperatorNotifier`: Crewing Manager (email), Vessel Master "
                "(email), the proposed crew (SMS), or, on a no-crew dead-end, an Ops-Center escalation.")
    bullet(doc, "**L4 (downstream).** The same ranked result + rationale + per-dimension scores is "
                "exactly the trace L4's Decision Graph would record as precedent.")

    h(doc, "4.5 The full L3 workflow (event timeline)", 2)
    para(doc, "One run emits this streamed sequence — the orchestrator's view, and what the "
              "live console renders:")
    code(doc, """10:01  intel_supervisor_started          (vacancy received, 3 investigators named)
10:01  intel_investigator_started  x3     (Crew / Vessel Ops / Contract-Wage, in parallel)
10:02  intel_investigator_completed x3    (each: N eligible / M assessed)
10:02  intel_ranking                      (top-3 + combined rationale)   --+ one or
       intel_no_crew                      (graceful dead-end)           --+ the other
10:02  intel_graph                        (derived fit graph, animates in live)
10:03  intel_notification_sent  xN        (operators notified on the right channel)
10:03  intel_supervisor_completed         (status, timing)""")

    # ── 5. Explainability ──────────────────────────────────────────────────────
    h(doc, "5. Explainability — Why #1, and Why the Fallback Scored Less", 1)
    para(doc, "Because `reports[]` carries every dimension's reasons + the facts consulted, the "
              "UI can show a full justification when a candidate is selected, instead of a bare "
              "score:")
    code(doc, """WHY #1 (Juan dela Cruz, 98.0)
  Crew Intel  (50%, +48.3)   v Available  v Exact rank (Chief Officer)  v STCW valid  v 12 yrs
  Vessel Ops  (30%, +30.0)   v Holds all mandated certs  v Meets sea-time  v Already at join port
  Contract    (20%, +20.0)   v Wage $8,525 within band $6,500-$9,000  v 6-month contract in envelope
  Sources: [L2] required certs, [L2] Singapore restrictions, wage band, port schedule

WHY #2 IS LOWER (Piotr Kowalski, 94.7)
  -3.6 pts on Vessel Ops  ->  #1 led with: "Already at join port (Singapore) - no relocation\"""")
    para(doc, "Each fact is labelled with its **source** (e.g. `[L2]` for graph-derived facts), so "
              "the reasoning ties back to the Knowledge Graph.")

    # ── 6. End-to-end walkthrough ──────────────────────────────────────────────
    h(doc, "6. End-to-End Walkthrough (one sign-off)", 1)
    numbered(doc, "**L1:** An operator clicks **Initiate Sign Off** on a departing Chief Officer. "
                  "The event is normalized into a `SignOffContext` (rank, vessel, port).")
    numbered(doc, "**L2:** L3 gathers the facts it needs — the candidate pool and the port's "
                  "restrictions / required certs — from the Knowledge Graph (or its rule data).")
    numbered(doc, "**L3 (processing):** The Supervisor delegates to Crew Intel, Vessel Ops and "
                  "Contract/Wage in parallel; each scores every candidate; hard gates filter the "
                  "ineligible; fusion ranks the rest.")
    numbered(doc, "**L3 (output):** A ranked top-3 with full rationale + the fit graph + operator "
                  "notifications. The UI auto-switches to the **Shortlist** tab.")
    numbered(doc, "**Action:** The agent signs on **#1**, moving them into the onboard pool; "
                  "**#2 / #3** remain as explained fallbacks.")
    numbered(doc, "**L4 (would-be):** the decision is handed off to be recorded as precedent.")
    note(doc, "**ROI:** manual sign-off ~8 hours / ~$250 per operation; automated ~2-5 minutes / "
              "~$0.01-0.05 — roughly 7.9 hours and ~$245 saved per operation.")

    # ── 7. Backends & data sources ─────────────────────────────────────────────
    h(doc, "7. Backends, Data Sources & the Latency SLO", 1)
    para(doc, "L3 has two independent, configurable axes — **how it reasons** and **where its "
              "data comes from** — so the same layer scales from a zero-cost demo to a full "
              "graph-backed, LLM-agent deployment without changing the reasoning or the UI.")
    table(doc, ["Axis / setting", "Options", "Effect"], [
        ["**Reasoning** — `INTEL_BACKEND`",
         "`fallback` (default) — deterministic Python; `managed` — Claude Managed-Agents "
         "coordinator + 3 LLM sub-agents",
         "fallback: ~ms, **zero tokens**, meets the <2s/<10s SLOs. managed: real LLM agents that "
         "delegate & narrate; scores still come from the deterministic tools, so ranking is identical"],
        ["**Data source** — `GRAPH_BACKEND`",
         "`fallback` (default) — rule data + relational pool; `age` — L2 Cypher + candidates "
         "from `(:Seafarer)` nodes",
         "Flips L3 from seed data to the live L2 knowledge graph. **Zero tokens** either way"],
    ])
    para(doc, "**Meeting the <10s SLO with real agents.** The managed coordinator + 3 sub-agents take "
              "60-90s end-to-end, so the response is served by the **fast deterministic path** "
              "(<10s) while the LLM agents run in the **background** and stream their reasoning as "
              "enrichment — the SLO is met and the real agents still contribute.")
    table(doc, ["Acceptance criterion", "Result"], [
        ["Supervisor delegates to all 3 investigators", "Pass"],
        ["Sign-off → top-3 ranked with rationale", "Pass"],
        ["First token < 2s, full response < 10s", "Pass (deterministic path; ms-scale)"],
        ["Crew notified via the correct channel", "Pass (email / SMS / Slack by role)"],
        ["5 sign-off scenarios pass", "Pass (6/6 in `verify_l3_intelligence`)"],
        ["“No crew found” handled gracefully", "Pass (escalates instead of proposing)"],
        ["Top-1 accuracy", "100% on matched scenarios"],
    ])

    # ── 8. Code map ────────────────────────────────────────────────────────────
    h(doc, "8. Where It Lives in the Code", 1)
    table(doc, ["Concern", "File(s)"], [
        ["Supervisor (orchestration)", "`backend/agents/intelligence/supervisor.py`"],
        ["3 investigators", "`crew_intel.py`, `vessel_ops_intel.py`, `contract_wage_intel.py`"],
        ["Ranking / fusion (the decider)", "`backend/agents/intelligence/ranking.py`"],
        ["Fit graph (visualization)", "`backend/agents/intelligence/fit_graph.py`"],
        ["Notifications", "`backend/agents/intelligence/notifications.py`"],
        ["Output shapes", "`backend/agents/intelligence/schemas.py`"],
        ["L2 seam + rule data", "`graph_gateway.py`, `database/intel_rules.py`, `database/graph_db.py`"],
        ["Graph-sourced candidates", "`database/crew_graph.py`"],
        ["Managed-agents variant", "`managed_supervisor.py`, `managed_registry.py`"],
        ["API entry points", "`backend/api/routes/intelligence.py`"],
        ["UI (shortlist / reasoning / graph)", "`frontend/src/components/{dashboard/ShortlistTab,intelligence/*}.tsx`"],
    ])

    para(doc, "")
    foot = doc.add_paragraph()
    fr = foot.add_run("Generated from the Maritime Crew Orchestrator codebase — L3 Intelligence Graph workflow.")
    fr.italic = True
    fr.font.size = Pt(8.5)
    fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
