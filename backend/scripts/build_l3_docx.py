"""
One-off generator: consolidate the two L3 markdown design docs
(L3_INTELLIGENCE_GRAPH.md + L3_FIT_GRAPH_DESIGN.md) into a single Word document.

Run:  python backend/scripts/build_l3_docx.py
Out:  docs/L3_Intelligence_Graph_Design.docx
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parents[2] / "docs" / "L3_Intelligence_Graph_Design.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
CODE_FILL = "F4F4F4"
INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def _shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _rich(p, text):
    """Render inline **bold** and `code` markdown into runs."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        else:
            p.add_run(part)


def para(doc, text="", style=None):
    p = doc.add_paragraph(style=style)
    _rich(p, text)
    return p


def bullet(doc, text):
    return para(doc, text, style="List Bullet")


def numbered(doc, text):
    return para(doc, text, style="List Number")


def code(doc, text):
    lines = text.strip("\n").split("\n")
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        _shade(p, CODE_FILL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        _rich(cell.paragraphs[0], h)
        for r in cell.paragraphs[0].runs:
            r.bold = True
        _shade(cell.paragraphs[0], "D9E2F3")
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(row):
            cells[i].paragraphs[0].text = ""
            _rich(cells[i].paragraphs[0], c)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def h(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level <= 1:
        run.font.color.rgb = ACCENT
    return heading


def build():
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ── Title block ───────────────────────────────────────────────────────────
    title = doc.add_heading(level=0)
    tr = title.add_run("L3 — Intelligence Graph")
    tr.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sr = sub.add_run("Consolidated Design Document — Architecture + Derived Fit Graph")
    sr.italic = True
    sr.font.size = Pt(12)

    meta = doc.add_paragraph()
    _rich(meta, "**Layer:** L3 Intelligence Graph    **Engineers:** Satish, Venny    "
                "**Date:** 2026-06-05    **Status:** Implemented · backend logic verified · frontend type-checked")
    para(doc, "Related: L3_TEST_PLAN.md · backend/scripts/verify_l3_intelligence.py")

    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Pt(18)
    qrun = quote.add_run(
        "Scope (from the build plan): Supervisor + specialist investigators for the Maritime "
        "Crew domain. Crew Intel — availability, certs, rank eligibility. Contract/Wage Intel — "
        "applicable rules for the period. Vessel Ops Intel — requirements and port schedule. "
        "Supervisor orchestrates all three, triggers match + notify operators. Streaming via "
        "Vinu. Notifications via Venny."
    )
    qrun.italic = True
    qrun.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ════════════════════════════════════════════════════════════════════════
    # PART A — L3 ARCHITECTURE
    # ════════════════════════════════════════════════════════════════════════
    h(doc, "Part A — L3 Architecture", 1)

    h(doc, "1. Purpose & position in the stack", 2)
    para(doc, "L3 is the **reasoning layer**. It consumes the entities/relationships the **L2 "
              "Knowledge Graph** exposes (EntityMap / OpsMap / OrgMap) and turns a **sign-off "
              "vacancy** into a **ranked, explained shortlist of replacement crew**, then notifies "
              "the operators who act on it.")
    code(doc,
         " L1 SignalFabric ─▶ L2 Knowledge Graph ─▶  L3 Intelligence Graph  ─▶ operators\n"
         " (event streams)    (EntityMap/OpsMap/      Supervisor + 3 investigators   (notify)\n"
         "                     OrgMap in PG+AGE)       └─ ranked top-3 + rationale\n"
         "                                                       ▲\n"
         "                                            L4 Decision Graph feeds precedent back")
    para(doc, "**Prototype stance (important):** L3 ships in the repo's **\"fallback\" mode** — the "
              "investigators are deterministic Python over rule data (`database/intel_rules.py`), so "
              "L3 is demoable and testable **today** with no API key, no graph infra, and no "
              "dependency on L2 being finished. The investigator interface is the seam: each can "
              "later be backed by a Managed-Agents sub-agent issuing Cypher against L2's graph "
              "without changing the Supervisor, the API, or the UI. This mirrors how "
              "`database/compliance_graph.py` already de-risks AGE for the Compliance Agent.")

    h(doc, "2. Supervisor–investigator pattern", 2)
    para(doc, "A **Supervisor** fans out to **three independent investigators**, each an expert in "
              "one dimension, then **fuses** their verdicts. Investigators never see each other and "
              "never rank across dimensions — that separation keeps each one independently testable "
              "and swappable.")
    code(doc,
         "Sign-off context ─▶ Supervisor\n"
         "                     ├─ Crew Intel          (availability · certs · rank eligibility)\n"
         "                     ├─ Contract/Wage Intel (applicable rules for the period)   [parallel]\n"
         "                     └─ Vessel Ops Intel     (requirements + port schedule)\n"
         "                     ▼\n"
         "                 fuse → top-3 ranked candidates (with rationale)\n"
         "                     ▼\n"
         "                 notify operators via the correct channel")
    para(doc, "Code map (`backend/agents/intelligence/`):")
    table(doc, ["File", "Responsibility"], [
        ["`supervisor.py`", "`IntelligenceSupervisor` — load pool, run 3 investigators in parallel, fuse, notify, stream, time"],
        ["`base_investigator.py`", "`BaseInvestigator` — per-candidate `_assess` contract + timing/event wrapping"],
        ["`crew_intel.py` / `contract_wage_intel.py` / `vessel_ops_intel.py`", "the three investigators"],
        ["`ranking.py`", "`fuse()` — hard gates + weighted blend → top-N with rationale"],
        ["`fit_graph.py`", "`build_fit_graph()` — derive the run's node/edge graph (vacancy→candidates→dimensions→L2 facts)"],
        ["`notifications.py`", "`OperatorNotifier` — channel routing + delivery (Venny's slice)"],
        ["`schemas.py`", "dataclasses (`SignOffContext`, `Assessment`, `RankedCandidate`, `IntelResult`, …)"],
        ["`database/intel_rules.py`", "wage bands, contract envelope, vessel certs, port schedule (fallback rule data)"],
        ["`api/routes/intelligence.py`", "`POST /api/v1/intelligence/match` and `/match-context`"],
    ])

    h(doc, "3. Investigator specs (inputs / outputs)", 2)
    para(doc, "All three implement `investigate(context, candidates) -> InvestigatorReport`, where "
              "the report holds one **`Assessment`** per candidate:")
    code(doc,
         "Assessment{ investigator, crew_id, score: 0..1, eligible: bool (hard gate),\n"
         "            signals: {...facts...}, reasons: [str] }")
    para(doc, "**Shared input — `SignOffContext`:** `vacated_rank`, `vacated_grade`, `vessel`, "
              "`port`, `sign_off_date`, `contract_period_months` (default 6), `workflow_id`.")

    h(doc, "3.1 Crew Intel — availability · certs · rank eligibility", 3)
    table(doc, ["Aspect", "Detail"], [
        ["Reads", "candidate `status`/`availability`, `rank`, `stcw_status`, `certifications`, `experience_years`"],
        ["Hard gates", "unavailable → ineligible; rank wrong family or >1 step on the rank ladder → ineligible"],
        ["Score (0..1)", "exact rank +0.55 / adjacent +0.35; STCW valid +0.25 (expiring +0.10); base certs +0.10; experience up to +0.10"],
        ["Signals", "`available`, `rank_distance`, `stcw_status`, `missing_base_certs`, `experience_years`"],
    ])
    para(doc, "Rank ladders (`intel_rules.py`): deck `Deck Cadet→Third→Second→Chief Officer→Master`, "
              "engine `Engine Cadet→Fourth→Third→Second→Chief Engineer`, ratings grouped. \"Adjacent\" "
              "(distance 1) is acceptable cover; cross-family is not.")

    h(doc, "3.2 Contract/Wage Intel — applicable rules for the period", 3)
    table(doc, ["Aspect", "Detail"], [
        ["Reads", "candidate `grade` (→ modelled wage), vacancy `vacated_rank`, `contract_period_months`"],
        ["Hard gates", "none (advisory dimension — surfaces commercial fit, never blocks alone)"],
        ["Score (0..1)", "wage in band +0.6 (below +0.5, over scaled down); contract within standard envelope +0.4 (within MLC max +0.25)"],
        ["Signals / applied", "`expected_wage_usd`, `wage_band_usd`, `period_months`, `contract_rules` (MLC-aligned)"],
    ])
    para(doc, "Modelled wage = band midpoint × grade premium (`GRADE_MULTIPLIER`). Envelope = "
              "`STANDARD_CONTRACT` (min 4 / max 9 / MLC max 11 months).")

    h(doc, "3.3 Vessel Ops Intel — requirements + port schedule", 3)
    table(doc, ["Aspect", "Detail"], [
        ["Reads", "candidate `certifications`, `experience_years`, `port`; vacancy `vacated_rank`, `port`"],
        ["Hard gates", "missing a **vessel-mandated cert** for the rank → ineligible"],
        ["Score (0..1)", "holds mandated certs +0.45; meets sea-time minimum +0.25; at join port +0.30 / can relocate in window +0.18"],
        ["Signals / applied", "`required_certs`, `missing_required_certs`, `min_experience_years`, `join_port`, `join_by`, `departure_window_days`"],
    ])
    para(doc, "Port schedule (`PORT_DEPARTURE_DAYS`) gives the **join-by** date; relocation needs "
              "`RELOCATION_LEAD_DAYS` of slack.")

    h(doc, "4. Fusion → top-3 with rationale", 2)
    para(doc, "`ranking.fuse(reports, candidates_by_id, top_n=3)`:")
    numbered(doc, "**Hard gate:** if **any** investigator marks a candidate `eligible=False`, the candidate is **disqualified** (and the gate reason is captured for \"why not\").")
    numbered(doc, "**Blend:** `score = 0.50·crew + 0.30·vessel + 0.20·contract` (×100).")
    numbered(doc, "**Rationale:** the top reason from each surviving dimension, concatenated — so an operator sees why a candidate placed where they did.")
    numbered(doc, "**Order:** score desc, `crew_id` asc as a stable tiebreak → assign `rank_position`, take top-N. Deterministic (no randomness) so scenarios have stable expected output.")
    para(doc, "Worked example (verified): CO vacancy at Singapore, 6-candidate pool → 4 disqualified "
              "by gates (missing GMDSS, unavailable, wrong family), shortlist = "
              "#1 Juan dela Cruz 98.0, #2 Rajesh Kumar 83.4 (adjacent-rank Second Officer).")

    h(doc, "5. Operator flows & notifications", 2)
    para(doc, "After fusion the **Operator Notifier** (`notifications.py`) dispatches to each "
              "recipient on their **correct channel** (`CHANNEL_BY_ROLE`):")
    table(doc, ["Recipient role", "Channel", "When"], [
        ["Crewing Manager", "email", "always (shortlist, or no-crew escalation)"],
        ["Vessel Master", "email", "on match (proposed candidate)"],
        ["Crew (proposed)", "**sms**", "on match (confirm availability)"],
        ["Ops Center", "slack", "on no-crew (manual-action alert)"],
    ])
    para(doc, "Delivery uses a mock sink (returns a delivery record) so L3 is testable without "
              "SMTP/Slack; the **channel-selection logic is the real, asserted behaviour**. The seam "
              "matches the existing `NotificationAgent`, so production swaps `_deliver` for the real "
              "mailer/Slack client without touching routing.")
    para(doc, "**No-crew-found flow:** when the shortlist is empty the Supervisor sets "
              "`status=\"no_crew_found\"`, emits `intel_no_crew`, and the Notifier escalates to "
              "Crewing Manager + Ops Center instead of proposing anyone — a graceful, audited dead-end.")

    h(doc, "6. Streaming architecture", 2)
    para(doc, "The Supervisor emits events through the **same callback → WebSocket vocabulary** the "
              "rest of the app uses, so the existing streaming layer / console render L3 with no "
              "change. Event sequence per run:")
    code(doc,
         "intel_supervisor_started\n"
         "  intel_investigator_started   × 3   (Crew / Contract-Wage / Vessel-Ops, in parallel)\n"
         "  intel_investigator_completed × 3\n"
         "intel_ranking            (top-N + rationale)         ─┐ one or the other\n"
         "intel_no_crew            (graceful dead-end)          ─┘\n"
         "intel_graph              (derived node/edge fit graph — animates in live)\n"
         "intel_notification_sent  × N\n"
         "intel_supervisor_completed   (status, timing)")
    para(doc, "**Latency design for the <2s first-token / <10s full SLOs:** the first event fires "
              "before any blocking work, and the deterministic Python path completes in single-digit "
              "milliseconds (measured `first_event_ms` and `total_ms` are returned in "
              "`IntelResult.timing`). When investigators are later backed by LLM sub-agents, the "
              "streaming-first ordering preserves the first-token budget while heavier reasoning runs "
              "behind it.")

    h(doc, "7. API", 2)
    code(doc,
         "POST /api/v1/intelligence/match\n"
         "  { \"crew_id\": \"<sign-off crew>\", \"contract_period_months\": 6, \"top_n\": 3 }\n"
         "  → IntelResult { status, candidates[], reports[], notifications[], timing, fit_graph, ... }\n"
         "\n"
         "POST /api/v1/intelligence/match-context     # explicit vacancy, no crew lookup\n"
         "  { \"vacated_rank\": \"Chief Officer\", \"port\": \"Singapore\", \"top_n\": 3 }")
    para(doc, "Events stream live on the global WebSocket (`/ws`).")

    # ════════════════════════════════════════════════════════════════════════
    # PART B — THE FIT GRAPH
    # ════════════════════════════════════════════════════════════════════════
    h(doc, "Part B — The Derived Fit Graph", 1)

    h(doc, "8. Background & motivation", 2)
    para(doc, "Before this change, \"L3 Intelligence **Graph**\" was a name, not a data structure. The "
              "layer computed everything you would need for a graph — a vacancy, a candidate pool, "
              "three scoring dimensions, scored/reasoned edges, and the L2 facts each assessment "
              "consulted — but **threw the structure away** into flat shapes (`candidates` list, "
              "`reports` dict, a scalar fused score). Three different things shared the word \"graph\" "
              "and only one was real:")
    table(doc, ["\"Graph\"", "What it actually was (before)"], [
        ["L3 \"Intelligence Graph\"", "A name for the supervisor→3-investigators→fusion topology"],
        ["Frontend `IntelligenceFlow.tsx`", "A hand-drawn flow diagram (fixed JSX nodes/connectors)"],
        ["L2 Knowledge/Compliance Graph (AGE)", "The only real property graph (nodes/edges/Cypher)"],
    ])
    para(doc, "We considered two senses of \"make L3 a graph\":")
    bullet(doc, "**Sense 1 — make L3's output a real graph.** Derive a node/edge graph from data the Supervisor already computes. No new infrastructure, no L2 dependency.")
    bullet(doc, "**Sense 2 — persist L3's graph into AGE.** Write placements into the same Apache AGE store L2 uses, so they become queryable precedent history. Requires AGE enabled and a write path.")
    para(doc, "**We implemented Sense 1**, plus a tiny streamed `intel_graph` event so the graph "
              "animates in live. Sense 1 is the highest value-for-effort: it makes the name literal, "
              "produces exactly the structure Sense 2 and L4 would later persist, and ships with zero "
              "new infra.")

    h(doc, "8.1 Goals / non-goals", 3)
    para(doc, "**Goals**")
    bullet(doc, "Turn each Supervisor run into a real, renderable node/edge graph.")
    bullet(doc, "Stream it live so the UI \"draws itself\" as the run lands.")
    bullet(doc, "Reuse the existing compliance-graph rendering vocabulary (status colors, x/y layout).")
    bullet(doc, "Keep it a pure transform — no new dependencies, no L2 requirement, deterministic.")
    para(doc, "**Non-goals (this change)**")
    bullet(doc, "Persisting the graph to AGE / making it queryable (that is Sense 2 / L4).")
    bullet(doc, "Changing the Supervisor's orchestration, the investigators, or the fusion math.")
    bullet(doc, "Building a generic graph engine — this is a purpose-built fit graph.")

    h(doc, "9. What the graph represents", 2)
    para(doc, "One run = one graph. A four-layer, left-to-right layered structure:")
    code(doc,
         " (Vacancy) ──assess──▶ (Candidate) ──score──▶ (Dimension) ──L2──▶ (L2 Fact)\n"
         "                         │\n"
         "                         └── disqualified candidates link to the BLOCKING dimension,\n"
         "                             carrying the gate reason as the edge label")
    para(doc, "**Semantics**")
    bullet(doc, "**Vacancy** — the sign-off context (vacated rank, port, vessel). One node.")
    bullet(doc, "**Candidate** — each crew member assessed. Status encodes outcome: `ok` (green) = shortlisted (ranked top-N); `warn` (amber) = eligible but not in the top-N; `block` (red) = disqualified by a hard gate.")
    bullet(doc, "**Dimension** — the three investigator lenses (Crew Intel / Vessel Ops / Contract-Wage), each labeled with its fusion weight (50% / 30% / 20%).")
    bullet(doc, "**L2 Fact** — the facts the investigators read from L2 (or its fallback rule data): the join port's nationality restrictions, and the rank's required safety certs. These tie the reasoning back to the L2 graph.")
    para(doc, "**Edges**")
    bullet(doc, "`Vacancy → Candidate`: `shortlist #n` (ok) · `disqualified` (block) · `assessed` (warn).")
    bullet(doc, "`Candidate → Dimension` (ranked only): the per-dimension score (e.g. `90`), `ok` if ≥ 0.5 else `warn` — the candidate's scoring breakdown.")
    bullet(doc, "`Candidate → blocking Dimension` (disqualified only): the gate reason as the label (e.g. \"Missing vessel-mandated certs for Chief Officer: GMDSS\"), colored `block`.")
    bullet(doc, "`Dimension → L2 Fact`: `L2 RESTRICTS` / `L2 REQUIRES`.")
    para(doc, "This is intentionally the same node/edge shape and `ok/warn/block` status vocabulary as "
              "the existing Compliance Context Graph, so the proven React-Flow renderer style applies "
              "directly.")

    h(doc, "10. Architecture & data flow", 2)
    code(doc,
         "POST /api/v1/intelligence/match-context\n"
         "        │\n"
         "        ▼\n"
         "IntelligenceSupervisor.find_replacements(context)\n"
         "        │  load pool → run 3 investigators (parallel) → fuse → top-N\n"
         "        ▼\n"
         "build_fit_graph(context, candidates_by_id, reports, ranked, backend)   ◀── NEW\n"
         "        │  pure transform → {nodes, edges, backend, node_count, edge_count}\n"
         "        ├─▶ result.fit_graph = <graph>\n"
         "        └─▶ emit \"intel_graph\" {nodes, edges, backend, counts}          ◀── NEW\n"
         "        ▼\n"
         "WebSocket /ws  ──broadcast──▶  frontend store (handleWSEvent)\n"
         "        │                              │  intel.fitGraph = <graph>\n"
         "        │                              ▼\n"
         "        │                      IntelligenceGraph.tsx (React-Flow)        ◀── NEW\n"
         "        │                         staggered node entrance animation\n"
         "        ▼\n"
         "HTTP response: IntelResult.to_dict() including fit_graph                 ◀── reconciled")
    para(doc, "Two delivery paths, both wired:")
    numbered(doc, "**Live** — the `intel_graph` WebSocket event arrives mid-run; the store sets `intel.fitGraph` and the component renders/animates immediately.")
    numbered(doc, "**Authoritative** — the HTTP response's `IntelResult.fit_graph` reconciles the final state (the store keeps the live graph if the payload omits it).")
    para(doc, "The graph is built and emitted in **both** outcomes — `matched` and `no_crew_found` "
              "(where every candidate is a red `block` node) — so the dead-end is still explained "
              "visually.")

    h(doc, "11. Backend implementation", 2)
    h(doc, "11.1 fit_graph.py (new)", 3)
    para(doc, "The core builder. Key design points:")
    bullet(doc, "**Pure & deterministic.** Depends only on `ranking._key_for`, `schemas`, and the data passed in. No I/O, no clock, no randomness → stable output for the same run.")
    bullet(doc, "**Layered layout.** Four fixed columns (`x = 0 / 240 / 520 / 780`), `GAP = 90`. `_column_ys(count, max_rows)` vertically centers each column; React-Flow `fitView` handles final scaling. Nodes are draggable.")
    bullet(doc, "**Bounded.** Candidates capped at `_MAX_CANDIDATES = 12` (ranked first, then disqualified, then eligible-not-ranked); any overflow surfaced as `+N more` on the vacancy node, not silently dropped.")
    bullet(doc, "**Classification.** `_disqualifying(crew_id)` scans reports in `crew → vessel → contract` order and returns the first hard gate `(dimension_key, reason)`; this both flags the candidate red and picks the dimension its block edge points at.")
    bullet(doc, "**L2 fact extraction.** Port rules from the Vessel Ops report's `applied`; safety certs from the first Crew Intel assessment carrying `signals.l2_required_safety_certs`; backend label resolved from what an investigator actually recorded.")
    para(doc, "Return shape:")
    code(doc,
         "{\n"
         "  \"nodes\": [{ \"id\", \"type\", \"label\", \"sub\", \"status\", \"x\", \"y\" }, ...],\n"
         "  \"edges\": [{ \"id\", \"source\", \"target\", \"label\", \"status\" }, ...],\n"
         "  \"backend\": \"age\" | \"fallback\",\n"
         "  \"node_count\": int,\n"
         "  \"edge_count\": int,\n"
         "}")

    h(doc, "11.2 schemas.py", 3)
    para(doc, "`IntelResult` gained `fit_graph: Optional[Dict[str, Any]] = None`, included in "
              "`to_dict()` so the API returns it.")

    h(doc, "11.3 supervisor.py", 3)
    para(doc, "After `fuse()` and the `intel_ranking` / `intel_no_crew` emit, and before notifications:")
    code(doc,
         "result.fit_graph = build_fit_graph(\n"
         "    context, candidates_by_id, list(reports), ranked, backend=l2_backend()\n"
         ")\n"
         "await emit(\"intel_graph\", {\n"
         "    \"workflow_id\": context.workflow_id,\n"
         "    \"nodes\": result.fit_graph[\"nodes\"],\n"
         "    \"edges\": result.fit_graph[\"edges\"],\n"
         "    \"backend\": result.fit_graph[\"backend\"],\n"
         "    \"node_count\": result.fit_graph[\"node_count\"],\n"
         "    \"edge_count\": result.fit_graph[\"edge_count\"],\n"
         "})")
    para(doc, "`l2_backend` is imported from `graph_gateway.backend` — the same seam the investigators "
              "use, so the reported backend is consistent with the actual L2 reads.")

    h(doc, "12. Frontend implementation", 2)
    h(doc, "12.1 IntelligenceGraph.tsx (new)", 3)
    para(doc, "React-Flow 11 view (already a project dependency, used by `ComplianceGraph.tsx`):")
    bullet(doc, "Reads `intel.fitGraph` from the store. Renders nothing until a graph exists.")
    bullet(doc, "**Custom `IntelNode`** mirrors the compliance node: status-colored ring + type-colored left accent (Vacancy cyan / Candidate sky / Dimension violet / L2Fact slate).")
    bullet(doc, "**Live entrance animation.** Each node is a `motion.div` with a staggered `delay = index*0.06s` spring. A per-run generation stamp (`intel.startedAt`) is passed as `data.gen` and used as the motion key, so nodes remount and re-animate on every run — the graph \"draws itself.\"")
    bullet(doc, "Status-colored edges; `block` edges are animated (flowing dashes). Arrow markers, label backgrounds — same styling as the compliance graph.")
    bullet(doc, "Footer legend (Shortlisted / Assessed / Disqualified) + live `L2 backend: …` label.")

    h(doc, "12.2 Store, types, and panel", 3)
    bullet(doc, "**Store (`workflowStore.ts`):** `IntelRunState` gained `fitGraph`; reset on `startIntelRun`; `handleWSEvent` handles the `intel_graph` case; `setIntelResult` reconciles `result.fit_graph ?? s.intel.fitGraph`; `intelLabel` adds \"Fit graph built — N nodes, M edges\".")
    bullet(doc, "**Types (`types/index.ts`):** `IntelGraphNode`, `IntelGraphEdge`, `IntelFitGraph` (reusing `GraphStatus`), plus `IntelResult.fit_graph?` and `IntelRunState.fitGraph`.")
    bullet(doc, "**Panel (`IntelligencePanel.tsx`):** renders `<IntelligenceGraph />` directly under the workflow pipeline diagram, so a run shows: pipeline flow → derived fit graph → ranked results.")

    h(doc, "13. Files changed", 2)
    table(doc, ["File", "Change"], [
        ["`backend/agents/intelligence/fit_graph.py`", "new — `build_fit_graph()`"],
        ["`backend/agents/intelligence/supervisor.py`", "build graph, set `fit_graph`, emit `intel_graph`"],
        ["`backend/agents/intelligence/schemas.py`", "`IntelResult.fit_graph` + serialization"],
        ["`frontend/src/components/intelligence/IntelligenceGraph.tsx`", "new — React-Flow view + animation"],
        ["`frontend/src/components/intelligence/IntelligencePanel.tsx`", "mount the graph component"],
        ["`frontend/src/store/workflowStore.ts`", "handle `intel_graph`, store/reset/reconcile `fitGraph`, label"],
        ["`frontend/src/types/index.ts`", "fit-graph types + fields"],
        ["`docs/L3_INTELLIGENCE_GRAPH.md` / `docs/L3_FIT_GRAPH_DESIGN.md`", "documented the new file, event, and concept"],
    ])

    h(doc, "14. Design decisions & rationale", 2)
    bullet(doc, "**Derive, don't re-engineer.** The graph is a pure projection of existing run data, so the Supervisor/investigators/fusion are untouched and the change is low-risk and testable in isolation.")
    bullet(doc, "**Reuse the compliance graph contract.** Same `{id,type,label,sub,status,x,y}` nodes, same `ok/warn/block` vocabulary, same React-Flow patterns → minimal new surface area and visual consistency.")
    bullet(doc, "**Build it in the Supervisor, not the API.** Only the Supervisor has the full `candidates_by_id` (needed for names/ranks of disqualified candidates, which the flat `reports` don't carry).")
    bullet(doc, "**Stream + reconcile.** Live event for responsiveness; HTTP `fit_graph` for the authoritative final state — the same dual-path pattern the rest of the L3 run already uses.")
    bullet(doc, "**Bounded with a visible note.** Capping at 12 candidates keeps the graph readable; surfacing `+N more` avoids \"silent truncation reads as complete coverage.\"")
    bullet(doc, "**Backend label honesty.** The graph reports `age` vs `fallback` from the same gateway the investigators used, so the UI never implies a live graph query that didn't happen.")

    h(doc, "15. Verification", 2)
    bullet(doc, "**Backend logic** — `build_fit_graph` validated in isolation (bypassing the heavy `agents` package init) against a synthetic Chief Officer @ Singapore run exercising all four candidate states (ranked #1/#2, cert-disqualified, eligible-not-ranked) and an L2 port restriction. 8 assertions pass: vacancy/dimension/candidate/L2-fact nodes present, disqualified candidate links to `dim_vessel` with the gate reason, ranked candidate has 3 scored edges, `fact_port` shows the restriction (warn), `fact_certs` carries the L2 required certs, and all edge ids are unique.")
    bullet(doc, "**Frontend** — `tsc --noEmit` → exit 0.")
    para(doc, "**Not run here:** the full FastAPI server (this session's Python was 3.14 with project "
              "deps absent; the dev env is 3.12). End-to-end verification = run the backend + frontend "
              "dev servers and hit Run Match; the `intel_graph` event streams over `/ws` and the graph "
              "animates in.")

    h(doc, "16. Dependencies, risks & productionisation seam", 2)
    bullet(doc, "**Depends on L2** for the real graph (Contract + Certification nodes, vessel/port schedule). Mitigation: prototype runs on `intel_rules.py` fallback data; the investigator interface is unchanged when L2 lands.")
    bullet(doc, "**Depends on L1** for fresh data; not needed for the prototype (seeded pool).")
    bullet(doc, "**Feeds L4** — the `IntelResult` (ranked candidates + rationale + per-dimension scores + `fit_graph`) is exactly the trace L4's Decision Graph records; L4's Precedent Index will later feed back as an additional fusion signal.")
    para(doc, "**Open decisions for review:** fusion weights (currently 50/30/20); whether Contract/Wage "
              "should ever hard-gate (e.g. wage > X% over band); rank-adjacency policy for senior ranks "
              "(should CO auto-cover Master?).")

    h(doc, "17. Future work", 2)
    bullet(doc, "**Sense 2 — persist to AGE.** Write the fit graph into the AGE store as `(:Candidate)-[:ASSESSED_BY {score}]->(:Vacancy)` etc., enabling precedent queries (e.g. \"CO placements at Singapore scoring > 80\"). Requires `age_enabled()` and a write path.")
    bullet(doc, "**L4 feedback.** `IntelResult.fit_graph` is exactly the trace L4's Decision Graph can record; L4's Precedent Index could feed back as an additional fusion signal.")
    bullet(doc, "**Richer L2 facts.** Expand fact nodes (medical-validity windows, port departure schedule, contract envelope) as more L2 reads are surfaced as signals.")
    bullet(doc, "**Interaction.** Click a candidate to expand full per-dimension rationale; highlight the path from a fact to the candidates it gated.")

    h(doc, "18. Status", 2)
    para(doc, "Prototype implemented and verified — see L3_TEST_PLAN.md and "
              "`backend/scripts/verify_l3_intelligence.py` (21/21 checks pass, all exit criteria "
              "exercised). The derived fit graph is implemented, backend logic verified, and the "
              "frontend type-checks clean.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
